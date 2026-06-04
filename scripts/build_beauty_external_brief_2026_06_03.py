from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
DOCX = OUT_DIR / "Beauty_Market_Rebuild_External_Brief_2026_06_03.docx"
SCREENSHOT = ROOT / "docs" / "beauty_market_qa_1440_2026_06_03.png"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for w in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(w))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Beauty 市场页重构外发简报")
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Amazon US 增长情报中台 · 2026-06-03")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("64748B")


def add_key_value_table(doc):
    rows = [
        ("验收页面", "http://127.0.0.1:8787/pages/market/?l1=Beauty"),
        ("重构范围", "仅 Beauty 市场页；未铺开其他一级行业。"),
        ("核心目标", "把市场页从 BI 报表改成 B2B 增长情报中台。"),
        ("最终状态", "核心趋势 + 左侧类目入口表 + 右侧主视觉详情面板 + 底部玩家与信号占位。"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    set_table_width(table, [2200, 7160])
    for i, (k, v) in enumerate(rows):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        set_cell_shading(table.cell(i, 0), "F2F4F7")
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(10)
            if i == 0:
                table.cell(i, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("2563EB")


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def add_l2_table(doc):
    l2s = [
        "功效面部护肤",
        "身体/沐浴/除臭",
        "彩妆/卸妆",
        "香水/香氛",
        "口腔护理",
        "男士剃须/理容",
        "女性脱毛/IPL",
        "洗护/头皮/防脱",
        "造型工具/吹风",
        "美甲/手足护理",
        "美容工具/仪器",
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [900, 2600, 4860])
    headers = ["#", "标准二级行业", "拆分逻辑"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)

    explanations = {
        "功效面部护肤": "承接 K-Beauty、成分功效、敏感肌与医美平替内容入口。",
        "身体/沐浴/除臭": "从护肤大桶中拆出身体护理、沐浴、除臭等高复购日用品逻辑。",
        "彩妆/卸妆": "保留试色、妆效、卸妆等内容入口，避免继续混入护肤。",
        "香水/香氛": "单独识别香评、平替、高性价比香氛和中东香水变量。",
        "口腔护理": "按电动牙刷、水牙线、牙膏、美白和耗材复购逻辑拆出。",
        "男士剃须/理容": "独立识别剃须刀、理发器、理容工具和男士替换周期。",
        "女性脱毛/IPL": "单独承载 IPL、脱毛仪、waxing 和内容测评型增长。",
        "洗护/头皮/防脱": "区分洗发护发、头皮护理、防脱和发膜发油复购。",
        "造型工具/吹风": "识别吹风机、直发器、卷发棒等参数和测评驱动品类。",
        "美甲/手足护理": "承载高 CN 渗透、教程化、色系上新和套装化机会。",
        "美容工具/仪器": "覆盖 LED、微电流、AGE-R、美容仪和家庭护理替代。",
    }
    for i, name in enumerate(l2s, 1):
        row = table.add_row()
        row.cells[0].text = str(i)
        row.cells[1].text = name
        row.cells[2].text = explanations[name]
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)


def add_validation_table(doc):
    rows = [
        ("node scripts\\render_research_portal_pages_v0_3.js", "通过", "rendered research portal pages v0.3"),
        ("node scripts\\validate_portal_pages_v0_1.js", "通过", "market / players / products / leads inline_js_ok"),
        ("node scripts\\validate_intelligence_portal_contract_v0_1.js", "通过", "intelligence portal contract ok"),
        ("1440x900 宽屏 QA", "通过", "无横向溢出；左右中部同顶同底；右侧详情内部滚动。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_width(table, [4300, 1200, 3860])
    for cell, text in zip(table.rows[0].cells, ["检查项", "结果", "说明"]):
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(10)
    for command, result, note in rows:
        row = table.add_row()
        row.cells[0].text = command
        row.cells[1].text = result
        row.cells[2].text = note
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(9.5)
        row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string("166534")


def build():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    style_doc(doc)
    add_title(doc)

    doc.add_heading("1. Executive Summary", level=1)
    add_key_value_table(doc)
    doc.add_paragraph(
        "本轮交付把 Beauty 市场页从数据报表式展示，改成面向 B2B 增长判断的情报工作台。"
        "首屏现在能够回答：Beauty 为什么值得做、哪些二级类目值得下钻、以及中国玩家在哪些结构里有真实机会。"
    )

    doc.add_heading("2. What Changed", level=1)
    add_bullets(
        doc,
        [
            "顶部从“核心观点”改为“核心趋势”，保留 3 条短而有信息量的 Beauty 趋势。",
            "Beauty 被拆成 11 个标准二级行业，替代粗粒度的“护肤与个护”主导视角。",
            "类目机会排行变成左侧入口表，仅展示最多 10 行，服务于点击下钻。",
            "中部右侧改成主视觉详情面板，承载指标、趋势图、玩家格局、产品机会、增长信号和推荐动作。",
            "底部玩家模块改为“中国玩家机会判断”；底部增长信号概览留作占位。",
        ],
    )

    doc.add_heading("3. Beauty Standard L2 Taxonomy", level=1)
    doc.add_paragraph(
        "拆分基于处理后的 Amazon US 市场底表、玩家主类目 main_l3、品牌名称和 Beauty 原始 L3 语义进行权重分配。"
        "这不是 SKU 级重建，而是用于市场页首屏判断的标准化业务口径。"
    )
    add_l2_table(doc)

    doc.add_heading("4. Core Trends Now Shown On Page", level=1)
    add_numbered(
        doc,
        [
            "K-Beauty 与功效护肤接管内容入口：medicube、ANUA、Beauty of Joseon 把成分证据变成可传播的购买理由。",
            "设备型个护正在消费电子化：IPL、造型工具、口腔护理和美容仪更容易用参数、测评和价格带建立差异。",
            "中国玩家机会集中在可参数化、教程化、耗材化细分：美甲、造型工具、IPL、口腔护理和部分美容仪优先下钻。",
        ],
    )

    doc.add_heading("5. Validation", level=1)
    add_validation_table(doc)

    if SCREENSHOT.exists():
        doc.add_heading("6. 1440x900 Visual QA Screenshot", level=1)
        doc.add_paragraph("下图为宽屏验收截图，用于外发说明页面首屏布局。")
        doc.add_picture(str(SCREENSHOT), width=Inches(6.5))

    doc.add_heading("7. Remaining Limits", level=1)
    add_bullets(
        doc,
        [
            "Beauty 拆分仍基于处理后底表和玩家主类目推断，不等于 SKU/ASIN 全量索引。",
            "增长事件仍需 PR、新闻、展会、TikTok 或 Google Trends 做持续补证。",
            "产品机会仍是方向级，不冒充真实 SKU 事实。",
            "其他一级行业尚未按 Beauty 新标准重构。",
        ],
    )

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Beauty Market Rebuild · 2026-06-03").font.size = Pt(9)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    build()
